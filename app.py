"""BambiVault Flask application and HTTP endpoints."""

import csv
import base64
import io
import json
import os
import secrets
import hmac
import math
import logging
from logging.handlers import RotatingFileHandler
from datetime import datetime, timedelta
from functools import wraps
from ai_feedback import get_ai_recommendations
import joblib
import numpy as np
from flask import (Flask, render_template, request, jsonify,
                    session, redirect, url_for, Response, send_file)
from flask_migrate import Migrate, upgrade as migrate_upgrade
from werkzeug.security import generate_password_hash, check_password_hash
from sqlalchemy import func
import pyotp
import qrcode
import email_utils as eu
from models import db, PasswordLog, User, UserSession, AdminAccount, AdminOTP
from feature_extraction import extract_features, rule_based_strength
from breach import check_breach, get_cache_stats
from strengthen import suggest_stronger_variants
from vault_routes import vault_bp
from database_migrations import run_database_migrations
from security_utils import consume_rate_limit, clear_rate_limit


from config import config

SECRET_KEY     = config.SECRET_KEY
ADMIN_PASSWORD = config.ADMIN_PASSWORD
DATABASE_URL   = config.DATABASE_URL
FLASK_DEBUG    = config.FLASK_DEBUG


# LOGGING

os.makedirs('logs', exist_ok=True)

log_handler = RotatingFileHandler('logs/bambivault.log', maxBytes=1_000_000, backupCount=3)
log_handler.setFormatter(logging.Formatter(
    '%(asctime)s | %(levelname)s | %(message)s', datefmt='%Y-%m-%d %H:%M:%S'
))

logger = logging.getLogger('bambivault')
logger.setLevel(logging.INFO)
logger.addHandler(log_handler)


# APP & CONFIG

app = Flask(__name__)
app.secret_key = SECRET_KEY

app.config['SQLALCHEMY_DATABASE_URI']        = DATABASE_URL
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['SQLALCHEMY_ENGINE_OPTIONS'] = {
    'pool_pre_ping': True,
    'pool_recycle': 300,
}

# Production configuration enforces secure session cookies.
app.config['SESSION_COOKIE_SECURE']   = config.SESSION_COOKIE_SECURE
app.config['SESSION_COOKIE_HTTPONLY'] = config.SESSION_COOKIE_HTTPONLY
app.config['SESSION_COOKIE_SAMESITE'] = config.SESSION_COOKIE_SAMESITE

db.init_app(app)
migrate = Migrate(app, db, compare_type=True)

app.register_blueprint(vault_bp)


def _bootstrap_database() -> None:
    """Apply migrations and create the initial administrator if required."""
    migrations_dir = os.path.join(os.path.dirname(__file__), 'migrations')
    if os.path.isdir(migrations_dir):
        try:
            migrate_upgrade(directory=migrations_dir)
        except Exception as exc:
            logger.critical('Database migration failed: %s', exc)
            raise
    else:
        run_database_migrations(db)
    admin_account = db.session.get(AdminAccount, 1)
    if not admin_account:
        admin_account = AdminAccount(
            id=1,
            email=(os.environ.get('ADMIN_EMAIL') or config.SMTP_USER or '').strip().lower() or None,
            password_hash=generate_password_hash(ADMIN_PASSWORD),
        )
        db.session.add(admin_account)
        db.session.commit()
    elif not admin_account.email:
        configured_admin_email = (os.environ.get('ADMIN_EMAIL') or config.SMTP_USER or '').strip().lower()
        if configured_admin_email:
            admin_account.email = configured_admin_email
            db.session.commit()


@app.cli.command('bootstrap')
def bootstrap_database_command() -> None:
    """Prepare the database once before starting production workers."""
    _bootstrap_database()
    logger.info('Database bootstrap completed')


# Production runs `flask --app app bootstrap` as a release/pre-deploy
# command. Keeping this out of worker imports prevents migration races.
if not config.is_production:
    with app.app_context():
        _bootstrap_database()


def _csrf_token():
    token = session.get('_csrf_token')
    if not token:
        token = secrets.token_urlsafe(32)
        session['_csrf_token'] = token
    return token


app.jinja_env.globals['csrf_token'] = _csrf_token


@app.before_request
def verify_csrf():
    """Require a session-bound token for every state-changing browser request."""
    if request.method not in {'POST', 'PUT', 'PATCH', 'DELETE'}:
        _csrf_token()
        return None
    expected = session.get('_csrf_token', '')
    supplied = request.headers.get('X-CSRF-Token') or request.form.get('csrf_token', '')
    if not expected or not supplied or not hmac.compare_digest(expected, supplied):
        if request.path.startswith('/api/') or request.is_json:
            return jsonify({'error': 'Security token missing or invalid. Refresh the page and try again.'}), 400
        return render_template('auth_result.html', title='Request unavailable', message='Refresh the page and try again.', success=False), 400


# SECURITY HEADERS
# Applied to every response. HSTS is only meaningful over real HTTPS, so it's
# gated on production — sending it during local HTTP development can cause
# browsers to "remember" HTTPS-only for 127.0.0.1 and break local testing.

@app.after_request
def set_security_headers(response):
    response.headers['X-Content-Type-Options'] = 'nosniff'
    response.headers['X-Frame-Options']        = 'DENY'
    response.headers['Referrer-Policy']        = 'strict-origin-when-cross-origin'
    response.headers['Permissions-Policy']     = 'camera=(), microphone=(), geolocation=(), payment=()'
    response.headers['Content-Security-Policy'] = (
        "default-src 'self'; base-uri 'self'; object-src 'none'; frame-ancestors 'none'; "
        "form-action 'self'; img-src 'self' data:; connect-src 'self'; "
        "font-src 'self' https://fonts.gstatic.com; "
        "style-src 'self' 'unsafe-inline' https://fonts.googleapis.com; "
        "script-src 'self' 'unsafe-inline' https://cdn.tailwindcss.com https://unpkg.com https://cdn.jsdelivr.net"
    )
    if config.is_production:
        response.headers['Strict-Transport-Security'] = 'max-age=63072000; includeSubDomains'
    return response


def admin_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        if not session.get('is_admin'):
            return redirect(url_for('admin_login'))
        return f(*args, **kwargs)
    return decorated


def login_required(f):
    """Require a valid user session and reject revoked sessions."""
    @wraps(f)
    def decorated(*args, **kwargs):
        uid = session.get('user_id')
        login_sid = session.get('login_session_id')

        if not uid or not login_sid:
            if request.path.startswith('/api/'):
                return jsonify({'error': 'Please log in to view this.'}), 401
            return redirect(url_for('login'))

        row = UserSession.query.filter_by(session_token=login_sid, user_id=uid).first()
        if not row:
            # Clear stale state after session revocation.
            session.pop('user_id', None)
            session.pop('login_session_id', None)
            session.pop('vault_token', None)
            if request.path.startswith('/api/'):
                return jsonify({'error': 'This session has been signed out. Please log in again.'}), 401
            return redirect(url_for('login'))

        row.last_activity = datetime.utcnow()
        db.session.commit()

        return f(*args, **kwargs)
    return decorated


def current_user():
    uid = session.get('user_id')
    if not uid:
        return None
    return User.query.get(uid)


def parse_date_range():
    start_str = request.args.get('start', '').strip()
    end_str   = request.args.get('end', '').strip()

    start_dt = None
    end_dt   = None
    try:
        if start_str:
            start_dt = datetime.strptime(start_str, '%Y-%m-%d')
        if end_str:
            end_dt = datetime.strptime(end_str, '%Y-%m-%d') + timedelta(days=1, seconds=-1)
    except ValueError:
        pass

    return start_dt, end_dt


def apply_date_filter(query, start_dt, end_dt):
    if start_dt:
        query = query.filter(PasswordLog.submitted_at >= start_dt)
    if end_dt:
        query = query.filter(PasswordLog.submitted_at <= end_dt)
    return query


# RATE LIMITING — protects /analyse from abuse

RATE_LIMIT_MAX    = 10   # requests
RATE_LIMIT_WINDOW = 60   # seconds


def is_rate_limited(ip: str) -> bool:
    """Use a shared database counter so limits work across all workers."""
    return consume_rate_limit('password-analysis', ip or 'unknown',
                              RATE_LIMIT_MAX, RATE_LIMIT_WINDOW)


VALID_BREACH_RISKS = {'Safe', 'Low Risk', 'Moderate Risk', 'High Risk', 'Critical', 'Unknown'}


def validate_analysis_result(feats, score, label, confidence, breach_result):
    """Reject malformed model/provider output before it reaches the UI or database."""
    if not isinstance(score, int) or not 0 <= score < len(STRENGTH_LABELS):
        raise ValueError('Invalid strength score')
    if label != STRENGTH_LABELS[score]:
        raise ValueError('Strength label does not match score')
    numeric = (feats.get('entropy'), feats.get('length'), confidence, breach_result.get('risk_score'))
    if any(not isinstance(v, (int, float)) or not math.isfinite(float(v)) for v in numeric):
        raise ValueError('Non-finite numeric result')
    if feats['length'] < 1 or feats['length'] > 1024 or feats['entropy'] < 0:
        raise ValueError('Feature values outside accepted range')
    if not 0 <= float(confidence) <= 100:
        raise ValueError('Invalid model confidence')
    if breach_result.get('risk_label') not in VALID_BREACH_RISKS:
        raise ValueError('Invalid breach risk')
    count = breach_result.get('breach_count')
    if not isinstance(count, int) or count < 0:
        raise ValueError('Invalid breach count')


# ML MODEL

from ml_classifier import classify_strength, STRENGTH_LABELS


def classify_behaviour_profile(feats: dict) -> str:
    if feats['has_keyboard_walk']:
        return 'Keyboard-Walk Type'
    if feats['has_year']:
        return 'Name+Year Type'
    if feats['has_common_sub']:
        return 'Substitution Type'
    if feats['has_dict_word']:
        return 'Dictionary-Word Type'
    return 'Clean'


def build_recommendations(feats: dict, score: int, breach_result: dict, profile: str) -> list[str]:
    recs = []
    profile_advice = {
        'Keyboard-Walk Type': "Your password shows reliance on keyboard sequences "
                               "(like 'qwerty' or 'asdf') — these are the first patterns "
                               "tested in any dictionary attack. Try a random passphrase instead.",
        'Name+Year Type':     "Your password follows a name-plus-year pattern — highly "
                               "predictable to attackers who know basic details about you. "
                               "Avoid birth years and personal dates entirely.",
        'Substitution Type':  "Your password uses common character substitutions (like '@' "
                               "for 'a') — attackers' dictionaries already account for these "
                               "swaps, so they add little real protection.",
        'Dictionary-Word Type': "Your password is built around a common dictionary word — "
                                 "even with added numbers or symbols, this remains a predictable "
                                 "starting point for attackers.",
    }
    if profile in profile_advice:
        recs.append(profile_advice[profile])
    if feats['length'] < 8:
        recs.append("Increase your password length to at least 12 characters — "
                     "length is the single biggest factor in password strength.")
    if feats['num_special'] == 0:
        recs.append("Add special characters such as !, @, #, or $ to significantly "
                     "increase unpredictability.")
    if breach_result['is_breached']:
        recs.append(f"This password was found in known breach databases "
                     f"{breach_result['breach_count']:,} times — it must be "
                     f"changed immediately.")
    if score < 3:
        recs.append("Consider using a password manager to generate and store "
                     "high-entropy unique passwords.")
    if not recs:
        recs.append("Good password. Ensure it is unique across all your accounts "
                     "and not reused anywhere.")
    return recs


# PUBLIC PAGE ROUTES

@app.route('/')
def index():
    if session.get('user_id'):
        return redirect(url_for('dashboard'))
    return render_template('landing.html')


@app.route('/analyser')
def analyser_page():
    """Render the analyser for guests and authenticated users."""
    return render_template('index.html')


@app.route('/guest')
def guest():
    """Start an anonymous analyser session."""
    session['guest_mode'] = True
    return redirect(url_for('analyser_page'))


@app.route('/exit-guest')
def exit_guest():
    """End the anonymous analyser session."""
    session.pop('guest_mode', None)
    return redirect(url_for('index'))


@app.route('/login')
def login():
    if session.get('user_id'):
        return redirect(url_for('dashboard'))
    return render_template('login.html')


@app.route('/register')
def register():
    if session.get('user_id'):
        return redirect(url_for('dashboard'))
    return render_template('register.html')


@app.route('/logout')
def logout():
    """End the user, vault, and tracked login sessions."""
    login_sid = session.pop('login_session_id', None)
    if login_sid:
        UserSession.query.filter_by(session_token=login_sid).delete()
        db.session.commit()
    token = session.pop('vault_token', None)
    if token:
        import vault_crypto as vc
        vc.clear_vault_key(token)
    session.pop('user_id', None)
    session.pop('guest_mode', None)
    return redirect(url_for('index'))


@app.route('/dashboard')
@login_required
def dashboard():
    return render_template('dashboard.html')


@app.route('/database')
@login_required
def database():
    return render_template('database.html')


@app.route('/vault')
@login_required
def vault_page():
    """Render the authenticated vault workspace."""
    return render_template('vault.html')


@app.route('/health')
def health():
    """Report application and database readiness."""
    db_ok = True
    try:
        db.session.execute(db.text('SELECT 1'))
    except Exception:
        db_ok = False

    status_code = 200 if db_ok else 503
    return jsonify({
        'status':    'ok' if db_ok else 'degraded',
        'database':  'ok' if db_ok else 'unreachable',
        'time':      datetime.utcnow().isoformat()
    }), status_code


# ADMIN AUTH ROUTES

@app.route('/admin/login', methods=['GET', 'POST'])
def admin_login():
    error = None
    if request.method == 'POST':
        ip = request.remote_addr or 'unknown'
        if consume_rate_limit('admin-login', ip, 5, 15 * 60):
            return render_template('admin_login.html', error='Too many sign-in attempts. Try again in 15 minutes.'), 429
        pw = request.form.get('password', '')
        admin_account = db.session.get(AdminAccount, 1)
        if admin_account and check_password_hash(admin_account.password_hash, pw):
            otp = db.session.get(AdminOTP, admin_account.id)
            if otp and otp.enabled:
                session.clear()
                session['admin_2fa_pending'] = admin_account.id
                return render_template('admin_login.html', require_2fa=True)
            # Admin and user workspaces are intentionally separate. Switching
            # to admin access clears any ordinary-user session in this browser
            # so user navigation and vault state cannot appear in admin pages.
            session.pop('user_id', None)
            session.pop('login_session_id', None)
            session.pop('vault_token', None)
            session['is_admin'] = True
            clear_rate_limit('admin-login', ip)
            logger.info(f"Admin login success | ip={request.remote_addr}")
            return redirect(url_for('admin'))
        error = 'Incorrect password.'
        logger.warning(f"Admin login FAILED | ip={request.remote_addr}")
    return render_template('admin_login.html', error=error)


@app.route('/admin/login/2fa', methods=['POST'])
def admin_login_2fa():
    ip = request.remote_addr or 'unknown'
    if consume_rate_limit('admin-2fa', ip, 5, 5 * 60):
        return render_template('admin_login.html', require_2fa=True, error='Too many verification attempts. Try again in five minutes.'), 429
    admin_id = session.get('admin_2fa_pending')
    otp = db.session.get(AdminOTP, admin_id) if admin_id else None
    token = request.form.get('token', '').strip()
    valid = bool(otp and otp.enabled and pyotp.TOTP(otp.secret).verify(token, valid_window=1))
    if not valid and otp and otp.recovery_codes:
        hashes = json.loads(otp.recovery_codes)
        match = next((h for h in hashes if check_password_hash(h, token.upper())), None)
        if match:
            hashes.remove(match)
            otp.recovery_codes = json.dumps(hashes)
            db.session.commit()
            valid = True
    if not valid:
        return render_template('admin_login.html', require_2fa=True, error='Incorrect authentication code.'), 401
    session.clear()
    session['is_admin'] = True
    clear_rate_limit('admin-2fa', ip)
    return redirect(url_for('admin'))


@app.route('/admin/logout')
def admin_logout():
    session.clear()
    return redirect(url_for('admin_login'))


@app.route('/admin/forgot-password', methods=['GET', 'POST'])
def admin_forgot_password():
    sent = False
    if request.method == 'POST':
        email = request.form.get('email', '').strip().lower()
        ip = request.remote_addr or 'unknown'
        if not consume_rate_limit('admin-recovery', ip, 5, 60 * 60):
            admin_account = AdminAccount.query.filter_by(email=email).first()
            if admin_account:
                eu.send_admin_reset_email(admin_account)
        sent = True
    return render_template('admin_forgot_password.html', sent=sent)


@app.route('/admin/reset-password/<token>', methods=['GET', 'POST'])
def admin_reset_password(token):
    admin_id, fingerprint, token_error = eu.parse_admin_reset_token(token)
    admin_account = db.session.get(AdminAccount, admin_id) if admin_id else None
    if token_error or not admin_account or fingerprint != eu._password_fingerprint(admin_account.password_hash):
        return render_template('auth_result.html', title='Reset link unavailable', message='This administrator reset link is invalid or has expired.', success=False), 400
    error = None
    if request.method == 'POST':
        password = request.form.get('password', '')
        confirm = request.form.get('confirm', '')
        if len(password) < 12:
            error = 'Use at least 12 characters for the administrator password.'
        elif password != confirm:
            error = 'Passwords do not match.'
        else:
            admin_account.password_hash = generate_password_hash(password)
            db.session.commit()
            session.clear()
            return redirect(url_for('admin_login'))
    return render_template('admin_reset_password.html', error=error)


@app.route('/admin/settings')
@admin_required
def admin_settings():
    admin_account = db.session.get(AdminAccount, 1)
    otp = db.session.get(AdminOTP, 1)
    return render_template('admin_settings.html', admin_email=admin_account.email or '', twofa_enabled=bool(otp and otp.enabled))


@app.route('/api/admin/change-password', methods=['POST'])
@admin_required
def admin_change_password():
    data = request.get_json(silent=True) or {}
    current = data.get('current_password', '')
    password = data.get('password', '')
    confirm = data.get('confirm', '')
    admin_account = db.session.get(AdminAccount, 1)
    if not check_password_hash(admin_account.password_hash, current):
        return jsonify({'error': 'Current password is incorrect.'}), 401
    if len(password) < 12:
        return jsonify({'error': 'Use at least 12 characters for the new password.'}), 400
    if password != confirm:
        return jsonify({'error': 'Passwords do not match.'}), 400
    admin_account.password_hash = generate_password_hash(password)
    db.session.commit()
    session.clear()
    return jsonify({'status': 'ok'})


@app.route('/api/admin/2fa/setup', methods=['POST'])
@admin_required
def admin_2fa_setup():
    admin_account = db.session.get(AdminAccount, 1)
    otp = db.session.get(AdminOTP, 1)
    if otp and otp.enabled:
        return jsonify({'error': 'Two-factor authentication is already enabled.'}), 400
    secret = pyotp.random_base32()
    if otp:
        otp.secret = secret
    else:
        otp = AdminOTP(admin_id=1, secret=secret, enabled=False)
        db.session.add(otp)
    db.session.commit()
    uri = pyotp.TOTP(secret).provisioning_uri(name=admin_account.email or 'Administrator', issuer_name='BambiVault Admin')
    image = qrcode.make(uri)
    buf = io.BytesIO()
    image.save(buf, format='PNG')
    return jsonify({'secret': secret, 'qr_data': 'data:image/png;base64,' + base64.b64encode(buf.getvalue()).decode()})


@app.route('/api/admin/2fa/verify', methods=['POST'])
@admin_required
def admin_2fa_verify():
    otp = db.session.get(AdminOTP, 1)
    token = (request.get_json(silent=True) or {}).get('token', '').strip()
    if not otp or not pyotp.TOTP(otp.secret).verify(token, valid_window=1):
        return jsonify({'error': 'Incorrect authentication code.'}), 401
    codes = [secrets.token_hex(4).upper() for _ in range(10)]
    otp.enabled = True
    otp.recovery_codes = json.dumps([generate_password_hash(code) for code in codes])
    db.session.commit()
    return jsonify({'status': 'ok', 'recovery_codes': codes})


@app.route('/api/admin/2fa/disable', methods=['POST'])
@admin_required
def admin_2fa_disable():
    password = (request.get_json(silent=True) or {}).get('password', '')
    admin_account = db.session.get(AdminAccount, 1)
    if not check_password_hash(admin_account.password_hash, password):
        return jsonify({'error': 'Password is incorrect.'}), 401
    otp = db.session.get(AdminOTP, 1)
    if otp:
        db.session.delete(otp)
        db.session.commit()
    return jsonify({'status': 'ok'})


@app.route('/admin')
@admin_required
def admin():
    return render_template('admin.html')


# API ROUTES

@app.route('/analyse', methods=['POST'])
def analyse():
    ip = request.remote_addr or 'unknown'

    if is_rate_limited(ip):
        logger.warning(f"Rate limit hit | ip={ip}")
        return jsonify({'error': 'The request limit has been reached. Try again after 60 seconds.'}), 429

    data     = request.get_json(silent=True) or {}
    password = data.get('password', '').strip()

    if not password:
        return jsonify({'error': 'No password provided'}), 400
    if len(password) > 1024:
        return jsonify({'error': 'Password input is too long.'}), 400

    feats = extract_features(password)
    score, label, confidence = classify_strength(feats)
    profile = classify_behaviour_profile(feats)
    breach_result = check_breach(password)
    try:
        validate_analysis_result(feats, score, label, confidence, breach_result)
    except ValueError as exc:
        logger.error('Analysis result validation failed | error=%s', exc)
        return jsonify({'error': 'The analysis could not be validated. Please try again.'}), 500
    ai_result = get_ai_recommendations(
        ip=ip,
        label=label,
        score=score,
        profile=profile,
        flags=feats['flags'],
        entropy=feats['entropy'],
        length=feats['length'],
        breach_result=breach_result
    )

    if ai_result['recommendations']:
        recs = ai_result['recommendations']
    else:
        # Fallback to the rule-based system, capped to 3 for a consistent
        # frontend experience whether or not the AI call succeeded.
        recs = build_recommendations(feats, score, breach_result, profile)[:3]

    log = PasswordLog(
        hash_prefix       = breach_result['hash_prefix'],
        strength_label    = label,
        pattern_flags     = ', '.join(feats['flags']),
        behaviour_profile = profile,
        entropy           = feats['entropy'],
        breach_exposed    = breach_result['is_breached'],
        breach_risk       = breach_result['risk_label'],
        user_id           = session.get('user_id')  # None for guests
    )
    db.session.add(log)
    db.session.commit()

    logger.info(f"Analyse | ip={ip} | strength={label} | profile={profile} | "
                f"breached={breach_result['is_breached']} | api_status={breach_result['api_status']}")

    return jsonify({
        'strength_label':    label,
        'strength_score':    score,
        'model_confidence':  confidence,
        'behaviour_profile': profile,
        'entropy':           feats['entropy'],
        'features':          feats,
        'flags':             feats['flags'],
        'breach_found':      breach_result['is_breached'],
        'breach_count':      breach_result['breach_count'],
        'breach_risk':       breach_result['risk_label'],
        'breach_score':      breach_result['risk_score'],
        'breach_colour':     breach_result['risk_colour'],
        'breach_advice':     breach_result['risk_advice'],
        'api_status':        breach_result['api_status'],
        'breach_age':        breach_result.get('breach_age'),
        'recommendations':   recs,
        'ai_feedback_source': ai_result['source']
    })


@app.route('/strengthen', methods=['POST'])
def strengthen():
    ip = request.remote_addr or 'unknown'

    if is_rate_limited(ip):
        logger.warning(f"Rate limit hit (strengthen) | ip={ip}")
        return jsonify({'error': 'The request limit has been reached. Try again after 60 seconds.'}), 429

    data     = request.get_json(silent=True) or {}
    password = data.get('password', '').strip()

    if not password:
        return jsonify({'error': 'No password provided'}), 400
    if len(password) > 1024:
        return jsonify({'error': 'Password input is too long.'}), 400

    feats    = extract_features(password)
    variants = suggest_stronger_variants(password, feats)

    variants_report = []
    seen_variants = set()
    for v in variants:
        if not isinstance(v, str) or not v or v == password or v in seen_variants or len(v) > 256:
            continue
        seen_variants.add(v)
        v_feats = extract_features(v)
        v_score, v_label, v_confidence = classify_strength(v_feats)
        if v_label != STRENGTH_LABELS[v_score] or not 0 <= v_confidence <= 100:
            continue
        improvements = []
        if v_feats['length'] > feats['length']:
            improvements.append('greater length')
        if v_feats['entropy'] > feats['entropy']:
            improvements.append('higher character diversity')
        if feats.get('has_keyboard_walk') and not v_feats.get('has_keyboard_walk'):
            improvements.append('no detected keyboard sequence')
        reason = ', '.join(improvements[:2]) or 'a less predictable structure'
        variants_report.append({
            'password':         v,
            'entropy':          v_feats['entropy'],
            'strength_label':   v_label,
            'strength_score':   v_score,
            'model_confidence': v_confidence,
            'reason':           reason.capitalize(),
        })

    # Rank strongest-first so the top suggestion is always the best option,
    # not just whichever strategy happened to run first.
    variants_report.sort(key=lambda v: (v['strength_score'], v['entropy']), reverse=True)

    logger.info(f"Strengthen | ip={ip} | variants_generated={len(variants_report)}")

    return jsonify({'variants': variants_report})


@app.route('/api/stats')
@login_required
def api_stats():
    uid = session['user_id']
    base_query = PasswordLog.query.filter_by(user_id=uid)

    total    = base_query.count()
    breached = base_query.filter_by(breach_exposed=True).count()

    dist = db.session.query(
        PasswordLog.strength_label, func.count(PasswordLog.id)
    ).filter(PasswordLog.user_id == uid).group_by(PasswordLog.strength_label).all()

    patterns_raw   = db.session.query(PasswordLog.pattern_flags) \
                        .filter(PasswordLog.user_id == uid).all()
    pattern_counts = {}
    for row in patterns_raw:
        if row.pattern_flags:
            for flag in row.pattern_flags.split(', '):
                flag = flag.strip()
                if flag and flag != 'none_detected':
                    pattern_counts[flag] = pattern_counts.get(flag, 0) + 1

    risk_dist = db.session.query(
        PasswordLog.breach_risk, func.count(PasswordLog.id)
    ).filter(PasswordLog.user_id == uid).group_by(PasswordLog.breach_risk).all()

    avg_entropy = db.session.query(func.avg(PasswordLog.entropy)) \
                    .filter(PasswordLog.user_id == uid).scalar() or 0
    cache_info  = get_cache_stats()

    distribution = [{'label': r[0], 'count': r[1]} for r in dist]
    top_patterns = sorted(pattern_counts.items(), key=lambda x: x[1], reverse=True)[:6]
    overall_recommendations = _generate_user_overall_recommendations(
        total, breached, avg_entropy, distribution, top_patterns
    )

    return jsonify({
        'total':        total,
        'breached':     breached,
        'avg_entropy':  round(avg_entropy, 2),
        'breach_rate':  round((breached / total * 100), 1) if total else 0,
        'distribution': distribution,
        'top_patterns': top_patterns,
        'risk_dist':    [{'label': r[0], 'count': r[1]} for r in risk_dist],
        'cache_info':   cache_info,
        'overall_ai_recommendations': overall_recommendations,
    })


def _generate_user_overall_recommendations(total, breached, avg_entropy,
                                           distribution, top_patterns):
    """Generate safe overall guidance from non-reversible aggregate metrics."""
    if not total:
        return ["Analyse a password first so BambiVault can identify your overall security habits."]

    recommendations = []
    weak_count = sum(
        row['count'] for row in distribution
        if row['label'] in {'Very Weak', 'Weak'}
    )
    weak_rate = weak_count / total * 100

    if breached:
        recommendations.append(
            f"Replace the {breached} analysed password{'s' if breached != 1 else ''} "
            "found in known breaches, and do not reuse them on other accounts."
        )

    if top_patterns:
        pattern_name = top_patterns[0][0].replace('_', ' ')
        pattern_advice = {
            'keyboard walk': "avoid keyboard sequences and use unrelated words instead",
            'name year combo': "avoid names and years that can be guessed from personal information",
            'common substitution': "do not rely on predictable substitutions such as '@' for 'a'",
            'dict word': "replace single dictionary words with a longer, unique passphrase",
        }
        advice = pattern_advice.get(pattern_name, f"reduce repeated use of the {pattern_name} pattern")
        recommendations.append(
            f"Your most frequent weakness is {pattern_name}; {advice}."
        )

    if weak_rate >= 25:
        recommendations.append(
            f"{round(weak_rate)}% of your analysed passwords are weak or very weak; "
            "prioritise longer, unique passphrases and save them in the encrypted vault."
        )
    elif avg_entropy < 80:
        recommendations.append(
            f"Your average entropy is {avg_entropy:.1f} bits; increase password length "
            "and randomness to move closer to the 80-bit target."
        )

    if not recommendations:
        recommendations.append(
            "Your analysed passwords show healthy overall habits; keep every password "
            "unique and continue using two-factor authentication where available."
        )

    return recommendations[:3]

def _generate_institutional_insight(stats: dict) -> list[str]:
    """Summarise aggregated behavioural results for administrators."""
    total = stats['total']
    if not total:
        return ["No submissions in this range yet."]

    profile_map = {p['label']: p['count'] for p in stats['profile_dist']}
    dominant_profile = max(profile_map, key=profile_map.get) if profile_map else 'Clean'
    dominant_pct = round((profile_map.get(dominant_profile, 0) / total) * 100, 1)

    profile_lines = {
        'Keyboard-Walk Type': (
            f"{dominant_pct}% of submissions show keyboard-walk patterns "
            f"(e.g. 'qwerty', 'asdf'), which automated cracking tools test "
            f"early. Prioritise awareness material addressing this weakness."
        ),
        'Name+Year Type': (
            f"{dominant_pct}% follow a name-plus-year structure — highly "
            f"predictable when an attacker knows personal details. Include "
            f"this risk in student security-awareness briefings."
        ),
        'Substitution Type': (
            f"{dominant_pct}% rely on common character substitutions (e.g. "
            f"'@' for 'a') — modern cracking dictionaries already account "
            f"for these swaps. This behaviour may result in greater exposure than users expect."
        ),
        'Dictionary-Word Type': (
            f"{dominant_pct}% are built around a single dictionary word — "
            f"the most crackable pattern observed, and the top priority for "
            f"follow-up awareness efforts."
        ),
        'Clean': (
            f"{dominant_pct}% of submissions show no detected weak pattern — "
            f"this indicates fewer detected behavioural weaknesses."
        ),
    }

    insights = [profile_lines.get(
        dominant_profile, f"Dominant profile: {dominant_profile} ({dominant_pct}%)."
    )]

    if stats['breach_rate'] > 30:
        insights.append(
            f"{stats['breach_rate']}% of submissions matched a known "
            f"breach. Consider escalating this finding for institutional "
            f"review independently of behavioural profile."
        )
    elif stats['breach_rate'] > 0:
        insights.append(
            f"{stats['breach_rate']}% of submissions matched a known "
            f"breach database."
        )

    return insights

def _compute_admin_stats(start_dt, end_dt):
    """Shared by /api/admin/stats and /api/admin/compare — one source of truth."""
    base = apply_date_filter(PasswordLog.query, start_dt, end_dt)
    total    = base.count()
    breached = apply_date_filter(PasswordLog.query.filter_by(breach_exposed=True), start_dt, end_dt).count()

    dist_q = apply_date_filter(
        db.session.query(PasswordLog.strength_label, func.count(PasswordLog.id)),
        start_dt, end_dt
    ).group_by(PasswordLog.strength_label).all()

    risk_q = apply_date_filter(
        db.session.query(PasswordLog.breach_risk, func.count(PasswordLog.id)),
        start_dt, end_dt
    ).group_by(PasswordLog.breach_risk).all()

    profile_q = apply_date_filter(
        db.session.query(PasswordLog.behaviour_profile, func.count(PasswordLog.id)),
        start_dt, end_dt
    ).group_by(PasswordLog.behaviour_profile).all()

    avg_entropy_q = apply_date_filter(
        db.session.query(func.avg(PasswordLog.entropy)), start_dt, end_dt
    ).scalar() or 0

    return {
        'total':        total,
        'breached':     breached,
        'avg_entropy':  round(avg_entropy_q, 2),
        'breach_rate':  round((breached / total * 100), 1) if total else 0,
        'distribution': [{'label': r[0], 'count': r[1]} for r in dist_q],
        'risk_dist':    [{'label': r[0], 'count': r[1]} for r in risk_q],
        'profile_dist': [{'label': r[0] or 'Clean', 'count': r[1]} for r in profile_q],
    }


@app.route('/api/admin/stats')
@admin_required
def api_admin_stats():
    start_dt, end_dt = parse_date_range()
    stats = _compute_admin_stats(start_dt, end_dt)
    stats['institutional_insight'] = _generate_institutional_insight(stats)

    day_expr = func.strftime('%Y-%m-%d', PasswordLog.submitted_at)
    trend_q = apply_date_filter(
        db.session.query(
            day_expr.label('day'),
            func.count(PasswordLog.id).label('total'),
            func.sum(func.cast(PasswordLog.breach_exposed, db.Integer)).label('breached')
        ),
        start_dt, end_dt
    ).group_by('day').order_by('day').limit(60).all()

    stats['breach_trend']  = [{'day': r[0], 'total': r[1], 'breached': r[2] or 0} for r in trend_q]
    stats['cache_info']    = get_cache_stats()
    stats['range_applied'] = {'start': start_dt.strftime('%Y-%m-%d') if start_dt else None,
                               'end':   end_dt.strftime('%Y-%m-%d') if end_dt else None}
    return jsonify(stats)


@app.route('/api/admin/compare')
@admin_required
def api_admin_compare():
    """
    Week-over-week (or custom range vs the immediately preceding equal-length
    period) comparison. If ?start=&end= given, the "previous" window is the
    same number of days immediately before start.
    """
    start_dt, end_dt = parse_date_range()

    if not start_dt or not end_dt:
        # default: current window = last 7 days, previous = the 7 before that
        end_dt   = datetime.utcnow()
        start_dt = end_dt - timedelta(days=7)

    window_len = end_dt - start_dt
    prev_end   = start_dt - timedelta(seconds=1)
    prev_start = prev_end - window_len

    current  = _compute_admin_stats(start_dt, end_dt)
    previous = _compute_admin_stats(prev_start, prev_end)

    return jsonify({
        'current':  current,
        'previous': previous,
        'current_range':  {'start': start_dt.strftime('%Y-%m-%d'), 'end': end_dt.strftime('%Y-%m-%d')},
        'previous_range': {'start': prev_start.strftime('%Y-%m-%d'), 'end': prev_end.strftime('%Y-%m-%d')},
    })


def _gather_export_data(start_dt, end_dt):
    strength = apply_date_filter(
        db.session.query(PasswordLog.strength_label, func.count(PasswordLog.id)),
        start_dt, end_dt
    ).group_by(PasswordLog.strength_label).all()

    profile = apply_date_filter(
        db.session.query(PasswordLog.behaviour_profile, func.count(PasswordLog.id)),
        start_dt, end_dt
    ).group_by(PasswordLog.behaviour_profile).all()

    risk = apply_date_filter(
        db.session.query(PasswordLog.breach_risk, func.count(PasswordLog.id)),
        start_dt, end_dt
    ).group_by(PasswordLog.breach_risk).all()

    total    = apply_date_filter(PasswordLog.query, start_dt, end_dt).count()
    breached = apply_date_filter(PasswordLog.query.filter_by(breach_exposed=True), start_dt, end_dt).count()

    return {
        'total': total,
        'breached': breached,
        'strength': strength,
        'profile': profile,
        'risk': risk,
        'range': f"{start_dt.strftime('%Y-%m-%d') if start_dt else 'All time'} to "
                 f"{end_dt.strftime('%Y-%m-%d') if end_dt else 'present'}",
        'range_suffix': f"{start_dt.strftime('%Y%m%d') if start_dt else 'all'}_"
                        f"{end_dt.strftime('%Y%m%d') if end_dt else 'present'}"
    }


def _generate_chart_png(data) -> bytes:
    """Renders strength + risk distribution as a single PNG for PDF embedding."""
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt

    fig, axes = plt.subplots(1, 2, figsize=(10, 3.5))
    fig.patch.set_facecolor('white')

    s_labels = ['Very Weak', 'Weak', 'Medium', 'Strong', 'Very Strong']
    s_map = {l: c for l, c in data['strength']}
    s_vals = [s_map.get(l, 0) for l in s_labels]
    s_colors = ['#ef4444', '#f97316', '#f59e0b', '#10b981', '#00ffc8']
    axes[0].bar(s_labels, s_vals, color=s_colors)
    axes[0].set_title('Strength Distribution', fontsize=10)
    axes[0].tick_params(axis='x', rotation=30, labelsize=7)

    r_labels = ['Safe', 'Low Risk', 'Moderate Risk', 'High Risk', 'Critical']
    r_map = {l: c for l, c in data['risk']}
    r_vals = [r_map.get(l, 0) for l in r_labels]
    r_colors = ['#10b981', '#f59e0b', '#f97316', '#ef4444', '#a855f7']
    axes[1].bar(r_labels, r_vals, color=r_colors)
    axes[1].set_title('Breach Risk Distribution', fontsize=10)
    axes[1].tick_params(axis='x', rotation=30, labelsize=7)

    plt.tight_layout()
    buf = io.BytesIO()
    plt.savefig(buf, format='png', dpi=140)
    plt.close(fig)
    buf.seek(0)
    return buf.read()


@app.route('/api/admin/export')
@admin_required
def api_admin_export():
    fmt = request.args.get('format', 'csv').lower()
    start_dt, end_dt = parse_date_range()
    data = _gather_export_data(start_dt, end_dt)

    logger.info(f"Admin export | format={fmt} | range={data['range']}")

    if fmt == 'pdf':
        return _export_pdf(data)
    elif fmt == 'docx':
        return _export_docx(data)
    elif fmt == 'txt':
        return _export_txt(data)
    else:
        return _export_csv(data)


def _export_csv(data):
    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow(['BambiVault — Statistics Report'])
    writer.writerow(['Range', data['range']])
    writer.writerow(['Total submissions', data['total']])
    writer.writerow(['Breached submissions', data['breached']])
    writer.writerow([])
    writer.writerow(['Strength Label', 'Count'])
    writer.writerows(data['strength'])
    writer.writerow([])
    writer.writerow(['Behaviour Profile', 'Count'])
    writer.writerows(data['profile'])
    writer.writerow([])
    writer.writerow(['Breach Risk Level', 'Count'])
    writer.writerows(data['risk'])

    return Response(
        output.getvalue(), mimetype='text/csv',
        headers={'Content-Disposition': f'attachment; filename=bambivault_stats_{data["range_suffix"]}.csv'}
    )


def _export_txt(data):
    lines = [
        "BAMBIVAULT — STATISTICS REPORT",
        "=" * 45,
        f"Range: {data['range']}",
        f"Total submissions: {data['total']}",
        f"Breached submissions: {data['breached']}",
        "",
        "STRENGTH DISTRIBUTION",
        "-" * 25,
    ]
    for label, count in data['strength']:
        lines.append(f"  {label:<15} {count}")
    lines += ["", "BEHAVIOUR PROFILE DISTRIBUTION", "-" * 32]
    for label, count in data['profile']:
        lines.append(f"  {(label or 'Clean'):<20} {count}")
    lines += ["", "BREACH RISK DISTRIBUTION", "-" * 27]
    for label, count in data['risk']:
        lines.append(f"  {label:<15} {count}")

    text = "\n".join(lines)
    return Response(
        text, mimetype='text/plain',
        headers={'Content-Disposition': f'attachment; filename=bambivault_stats_{data["range_suffix"]}.txt'}
    )


def _export_pdf(data):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.units import cm
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image
    from reportlab.lib.styles import getSampleStyleSheet

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, topMargin=2*cm, bottomMargin=2*cm)
    styles = getSampleStyleSheet()
    story = [
        Paragraph("BambiVault — Statistics Report", styles['Title']),
        Paragraph(f"Range: {data['range']}", styles['Normal']),
        Paragraph(f"Total submissions: {data['total']} &nbsp;&nbsp; "
                  f"Breached: {data['breached']}", styles['Normal']),
        Spacer(1, 16),
    ]

    # Chart image — embedded so the PDF is a self-contained visual report
    try:
        chart_bytes = _generate_chart_png(data)
        chart_buf = io.BytesIO(chart_bytes)
        story.append(Image(chart_buf, width=16*cm, height=5.6*cm))
        story.append(Spacer(1, 14))
    except Exception as e:
        logger.warning(f"Chart embed failed in PDF export: {e}")

    def make_table(title, rows):
        story.append(Paragraph(title, styles['Heading2']))
        table_data = [['Label', 'Count']] + [[str(r[0] or 'Clean'), str(r[1])] for r in rows]
        t = Table(table_data, colWidths=[8*cm, 4*cm])
        t.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1a1a2e')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('FONTSIZE', (0, 0), (-1, -1), 9),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ]))
        story.append(t)
        story.append(Spacer(1, 14))

    make_table("Strength Distribution", data['strength'])
    make_table("Behaviour Profile Distribution", data['profile'])
    make_table("Breach Risk Distribution", data['risk'])

    doc.build(story)
    buf.seek(0)
    return send_file(buf, mimetype='application/pdf', as_attachment=True,
                      download_name=f'bambivault_stats_{data["range_suffix"]}.pdf')


def _export_docx(data):
    from docx import Document
    from docx.shared import Cm

    doc = Document()
    doc.add_heading('BambiVault — Statistics Report', level=1)
    doc.add_paragraph(f"Range: {data['range']}")
    doc.add_paragraph(f"Total submissions: {data['total']}    Breached: {data['breached']}")

    try:
        chart_bytes = _generate_chart_png(data)
        chart_buf = io.BytesIO(chart_bytes)
        doc.add_picture(chart_buf, width=Cm(16))
    except Exception as e:
        logger.warning(f"Chart embed failed in DOCX export: {e}")

    def make_table(title, rows):
        doc.add_heading(title, level=2)
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Light Grid Accent 1'
        hdr = table.rows[0].cells
        hdr[0].text, hdr[1].text = 'Label', 'Count'
        for label, count in rows:
            cells = table.add_row().cells
            cells[0].text = str(label or 'Clean')
            cells[1].text = str(count)
        doc.add_paragraph()

    make_table("Strength Distribution", data['strength'])
    make_table("Behaviour Profile Distribution", data['profile'])
    make_table("Breach Risk Distribution", data['risk'])

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return send_file(buf,
                      mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                      as_attachment=True, download_name=f'bambivault_stats_{data["range_suffix"]}.docx')


@app.route('/api/records')
@login_required
def api_records():
    page       = max(request.args.get('page', 1, type=int), 1)
    per_page   = max(1, min(request.args.get('per_page', 15, type=int), 100))
    strength_f = request.args.get('strength', '')
    breach_f   = request.args.get('breach', '')
    profile_f  = request.args.get('profile', '')

    if strength_f and strength_f not in STRENGTH_LABELS:
        return jsonify({'error': 'Invalid strength filter.'}), 400
    if breach_f not in {'', 'yes', 'no'}:
        return jsonify({'error': 'Invalid breach filter.'}), 400
    valid_profiles = {'', 'Clean', 'Keyboard-Walk Type', 'Name+Year Type', 'Substitution Type', 'Dictionary-Word Type'}
    if profile_f not in valid_profiles:
        return jsonify({'error': 'Invalid behavioural profile filter.'}), 400

    query = PasswordLog.query.filter_by(user_id=session['user_id'])
    if strength_f:
        query = query.filter_by(strength_label=strength_f)
    if breach_f == 'yes':
        query = query.filter_by(breach_exposed=True)
    elif breach_f == 'no':
        query = query.filter_by(breach_exposed=False)
    if profile_f:
        if profile_f == 'Clean':
            query = query.filter(
                (PasswordLog.behaviour_profile == 'Clean') |
                (PasswordLog.behaviour_profile.is_(None))
            )
        else:
            query = query.filter_by(behaviour_profile=profile_f)

    records = query.order_by(PasswordLog.submitted_at.desc()) \
                    .paginate(page=page, per_page=per_page, error_out=False)

    return jsonify({
        'records': [r.to_dict() for r in records.items],
        'total':   records.total,
        'pages':   records.pages,
        'page':    page
    })


if __name__ == '__main__':
    logger.info("BambiVault starting up")
    app.run(debug=FLASK_DEBUG)
